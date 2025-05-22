from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def add_centered_text(paragraph, text, bold=False, size=12):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run

def add_right_text(paragraph, text, bold=False, size=12):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return run

def add_left_text(paragraph, text, bold=False, size=12):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return run

def export_medicalrecord_to_word(record):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    # Logo
    if os.path.exists('static/images/logo.png'):
        doc.add_picture('static/images/logo.png', width=Cm(2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Header
    p = doc.add_paragraph()
    add_centered_text(p, "Y KHOA UNG BƯỚU CẦN THƠ", bold=True, size=14)
    p = doc.add_paragraph()
    add_centered_text(p, "PHÒNG KHÁM MEDIONCON", bold=True, size=14)
    p = doc.add_paragraph()
    add_centered_text(p, "Địa chỉ: Số 10, ĐS 5, Tổ 17, KV Bình Thường B, P. Long Tuyền, Q. Bình Thủy, TPCT.", size=12)
    p = doc.add_paragraph()
    add_centered_text(p, "Điện thoại: 0917.575656.", size=12)
    p = doc.add_paragraph()
    add_centered_text(p, "Wed: Ykhoaungbuoucantho.com.vn.", size=12)
    p = doc.add_paragraph()
    add_centered_text(p, "Email.CSKH.MIDEONCO@GMAIL.COM.", size=12)
    # Thông tin phiếu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, f"STT: {datetime.now().strftime('%d/%m/%Y')}-{record.id}", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, f"Ngày xuất file: {datetime.now().strftime('%d/%m/%Y - %H:%M')}", size=12)
    # Tiêu đề
    p = doc.add_paragraph()
    add_centered_text(p, "BỆNH ÁN", bold=True, size=14)
    doc.add_paragraph()
    # Thông tin bệnh nhân
    p = doc.add_paragraph()
    add_left_text(p, f"Họ tên: {record.patient.full_name}", size=12)
    p = doc.add_paragraph()
    add_left_text(p, f"Ngày khám: {record.record_date.strftime('%d/%m/%Y')}", size=12)
    # Chẩn đoán, ghi chú
    if record.diagnosis:
        p = doc.add_paragraph()
        add_left_text(p, f"Chẩn đoán: {record.diagnosis}", size=12)
    if record.notes:
        p = doc.add_paragraph()
        add_left_text(p, f"Ghi chú: {record.notes}", size=12)

    # Bảng kết quả xét nghiệm
    labtests = record.lab_tests.select_related('category').prefetch_related('results__item').order_by('category__name', 'created_at')
    if labtests.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_left_text(p, "KẾT QUẢ XÉT NGHIỆM", bold=True, size=12)
        
        # Tạo bảng với 4 cột
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Đặt header
        headers = ['XÉT NGHIỆM', 'KẾT QUẢ', 'KHOẢN THAM CHIẾU', 'ĐƠN VỊ']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Thêm dữ liệu cho từng loại xét nghiệm
        current_category = None
        for lt in labtests:
            # Nếu là danh mục mới, thêm dòng phân cách
            if current_category != lt.category.name:
                if current_category is not None:  # Không thêm dòng trống ở đầu
                    row_cells = table.add_row().cells
                    for cell in row_cells:
                        cell.merge(row_cells[0])
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                current_category = lt.category.name
                # Thêm tên danh mục
                row_cells = table.add_row().cells
                for cell in row_cells:
                    cell.merge(row_cells[0])
                cell = row_cells[0]
                cell.text = current_category
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Thêm các trường xét nghiệm
            for result in lt.results.all():
                row_cells = table.add_row().cells
                row_cells[0].text = result.item.name
                row_cells[1].text = str(result.value)
                row_cells[2].text = result.item.reference_range
                row_cells[3].text = result.item.unit
                
                # Căn giữa tất cả các ô
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()
    # Footer
    p = doc.add_paragraph()
    add_centered_text(p, "Bệnh án đã được kiểm duyệt /QA", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, f"TP. Cần Thơ, ngày {datetime.now().strftime('%d')} tháng {datetime.now().strftime('%m')} năm {datetime.now().strftime('%Y')}", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, "PHÒNG KHÁM", size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, "(Ký, đóng dấu và ghi rõ họ tên)", size=12)
    # Lưu file
    filename = f"medicalrecord_{record.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename 
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import barcode
from barcode.writer import ImageWriter
from io import BytesIO

# Mapping gender values to Vietnamese
GENDER_MAP = {
    'Male': 'Nam',
    'Female': 'Nữ',
    'Other': 'Khác'
}

def add_centered_text(paragraph, text, bold=False, size=12, color=None):
    """Thêm text được căn giữa vào paragraph"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color and isinstance(color, tuple) and len(color) == 3:
        run.font.color.rgb = RGBColor(*color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run

def add_right_text(paragraph, text, bold=False, size=12, color=None):
    """Thêm text được căn phải vào paragraph"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color and isinstance(color, tuple) and len(color) == 3:
        run.font.color.rgb = RGBColor(*color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return run

def add_left_text(paragraph, text, bold=False, size=12, color=None):
    """Thêm text được căn trái vào paragraph"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color and isinstance(color, tuple) and len(color) == 3:
        run.font.color.rgb = RGBColor(*color)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return run

def generate_barcode(data):
    """Generate Code128 barcode image"""
    barcode_instance = barcode.get('code128', data, writer=ImageWriter())
    buffer = BytesIO()
    barcode_instance.write(buffer)
    buffer.seek(0)
    return buffer

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                    bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                    start={"sz": 24, "val": "dashed", "shadow": "true"},
                    end={"sz": 12, "val": "dashed"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # List over all available tags
    for edge in ['start', 'top', 'end', 'bottom', 'insideH', 'insideV']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # Check for tag existance, if none found, then create one
            element = tcBorders.find(qn(tag)) or OxmlElement(tag)
            if element.getparent() is None:
                tcBorders.append(element)

            # Set default values if not provided
            if 'sz' not in edge_data:
                edge_data['sz'] = 0
            if 'val' not in edge_data:
                edge_data['val'] = 'single'
            if 'color' not in edge_data:
                edge_data['color'] = '#FFFFFF'

            # Set attributes
            for key, value in edge_data.items():
                try:
                    element.set(qn(key), str(value))
                except ValueError:
                    # If setting the attribute fails, skip it
                    continue

def remove_table_borders(table):
    """Remove all borders from every cell in a python-docx table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top    ={"sz": 0, "val": "none", "color": "FFFFFF"},
                bottom ={"sz": 0, "val": "none", "color": "FFFFFF"},
                start  ={"sz": 0, "val": "none", "color": "FFFFFF"},
                end    ={"sz": 0, "val": "none", "color": "FFFFFF"},
            )

def remove_table_level_borders(table):
    """Strip out the <w:tblBorders> element so no outer frame is drawn."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)

def whiten_table_borders(table, sz=4):
    """
    Paint all table‐level and cell‐level borders white.
    `sz` is the border thickness in eighths of a point (4 ≃ 0.5pt).
    """
    # 1) Table‐level borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders')) or OxmlElement('w:tblBorders')
    if tblBorders.getparent() is None:
        tblPr.append(tblBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = tblBorders.find(qn(f'w:{edge}')) or OxmlElement(f'w:{edge}')
        if el.getparent() is None:
            tblBorders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), 'FFFFFF')
        el.set(qn('w:space'), '0')

    # 2) Cell‐level borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top    ={"sz": sz, "val": "single", "color": "FFFFFF"},
                bottom ={"sz": sz, "val": "single", "color": "FFFFFF"},
                start  ={"sz": sz, "val": "single", "color": "FFFFFF"},
                end    ={"sz": sz, "val": "single", "color": "FFFFFF"},
            )

def export_labtest_to_word(labtest):
    """Xuất kết quả xét nghiệm ra file Word"""
    doc = Document()
    
    # Thiết lập margins (1 inch = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        # Set up footer
        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6.5))  # Footer table needs width parameter
        footer_table.style = 'Table Grid'  # Keep grid for alignment
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Column 1: Quality Standards
        cell = footer_table.cell(0, 0)
        p = cell.paragraphs[0]
        add_centered_text(p, "TIÊU CHUẨN & CHẤT LƯỢNG", bold=True, size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "* An toàn sinh học cấp 2 – ATSH Class II", size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "* ISO 9001 : 2015 – Định hướng ISO 15189 : 2022", size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "* Miễn dịch tự động – Sinh học phân tử – Tế bào học", size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "* ThS (Giảng viên) Y khoa TP. HCM quản lý chất lượng", size=12, color=(84, 101, 154))

        # Column 2: Turnaround Time
        cell = footer_table.cell(0, 1)
        p = cell.paragraphs[0]
        add_centered_text(p, "THỜI GIAN TRẢ KẾT QUẢ:", bold=True, size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "- 60 phút cho xét nghiệm thường qui", size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "- 120 phút cho xét nghiệm thường qui + miễn dịch", size=12, color=(84, 101, 154))
        p = cell.add_paragraph()
        add_centered_text(p, "- 180 phút cho xét nghiệm ký sinh trùng, giun sán", size=12, color=(84, 101, 154))

        # Column 3: Flags
        cell = footer_table.cell(0, 2)
        p = cell.paragraphs[0]
        add_centered_text(p, "H: High / Cao  L: Low / Thấp", size=12, color=(84, 101, 154))

        # Make all borders white
        whiten_table_borders(footer_table)

    # Header (3-column layout)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = 'Table Grid'  # Keep grid for alignment
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.allow_autofit = False
    header_table.columns[0].width = Inches(2.17)
    header_table.columns[1].width = Inches(2.17)
    header_table.columns[2].width = Inches(2.17)
    
    # Column 1: Logo
    if os.path.exists('static/images/logo.png'):
        cell = header_table.cell(0, 0)
        cell.paragraphs[0].add_run().add_picture('static/images/logo.png', width=Cm(2.5))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column 2: Contact Info
    cell = header_table.cell(0, 1)
    p = cell.paragraphs[0]
    add_left_text(p, "TRUNG TÂM XÉT NGHIỆM MEDICON CẦN THƠ", bold=True, size=14, color=(255, 0, 0))
    p = cell.add_paragraph()
    add_left_text(p, "Lab Center: 40 Tân Khai, P.4, Q.11, TP. Hồ Chí Minh", size=12, color=(84, 101, 154))
    p = cell.add_paragraph()
    add_left_text(p, "Web: www.medilab.vn - www.medilab.com.vn", size=12, color=(84, 101, 154))
    p = cell.add_paragraph()
    add_left_text(p, "Email: cskh@medilab.vn", size=12, color=(84, 101, 154))

    # Column 3: Barcode and IDs
    cell = header_table.cell(0, 2)
    p = cell.paragraphs[0]
    # Create a table for barcodes to ensure same height
    barcode_table = cell.add_table(rows=1, cols=2)
    barcode_table.style = 'Table Grid'
    barcode_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    barcode_table.autofit = False
    barcode_table.allow_autofit = False
    barcode_table.columns[0].width = Inches(1.1)
    barcode_table.columns[1].width = Inches(1.1)
    
    # Generate barcode for patient ID
    barcode_buffer = generate_barcode(labtest.patient.patient_code)
    barcode_cell = barcode_table.cell(0, 0)
    barcode_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    barcode_cell.paragraphs[0].add_run().add_picture(barcode_buffer, width=Cm(1.2), height=Cm(0.75))  # 40% width, 25% height
    
    # Format STT: DDMMYY-XXXX
    stt = f"{labtest.created_at.strftime('%d%m%y')}-{labtest.id:04d}"
    # Generate barcode for STT
    stt_barcode_buffer = generate_barcode(stt)
    stt_cell = barcode_table.cell(0, 1)
    stt_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    stt_cell.paragraphs[0].add_run().add_picture(stt_barcode_buffer, width=Cm(0.84), height=Cm(0.84))  # 28% width, 28% height
    
    # Make barcode table borders white
    whiten_table_borders(barcode_table)
    
    # Text lines
    p = cell.add_paragraph()
    add_left_text(p, f"ID: {labtest.patient.patient_code}", size=12, color=(84, 101, 154))
    p = cell.add_paragraph()
    add_left_text(p, f"STT: {stt}", size=12, color=(84, 101, 154))
    p = cell.add_paragraph()
    add_left_text(p, f"Ngày đăng ký: {labtest.created_at.strftime('%d/%m/%Y')}", size=12, color=(84, 101, 154))
    p = cell.add_paragraph()
    add_left_text(p, f"Giờ in kết quả: {datetime.now().strftime('%H:%M')}", size=12, color=(84, 101, 154))

    # Make all borders white
    whiten_table_borders(header_table)

    # Main title
    p = doc.add_paragraph()
    add_centered_text(p, "KẾT QUẢ XÉT NGHIỆM", bold=True, size=16, color=(0, 1, 76))

    # Patient Info (3-column layout)
    info_table = doc.add_table(rows=3, cols=3)
    info_table.style = 'Table Grid'  # Keep grid for alignment
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.allow_autofit = False
    info_table.columns[0].width = Inches(2.17)
    info_table.columns[1].width = Inches(2.17)
    info_table.columns[2].width = Inches(2.17)
    
    # Row 1: Name, Birth Year, Gender
    cell = info_table.cell(0, 0)
    add_left_text(cell.paragraphs[0], f"Họ tên: {labtest.patient.full_name or ''}", bold=True, size=12, color=(84, 101, 154))
    cell = info_table.cell(0, 1)
    add_left_text(cell.paragraphs[0], f"Năm sinh: {labtest.patient.date_of_birth.year if labtest.patient.date_of_birth else ''}", size=12, color=(84, 101, 154))
    cell = info_table.cell(0, 2)
    gender_vn = GENDER_MAP.get(labtest.patient.gender, labtest.patient.gender or '')
    add_left_text(cell.paragraphs[0], f"Giới tính: {gender_vn}", size=12, color=(84, 101, 154))
    
    # Row 2: Address, Phone
    cell = info_table.cell(1, 0)
    add_left_text(cell.paragraphs[0], f"Địa chỉ: {labtest.patient.address or ''}", size=12, color=(84, 101, 154))
    cell = info_table.cell(1, 1)
    add_left_text(cell.paragraphs[0], f"Điện thoại: {labtest.patient.phone or ''}", size=12, color=(84, 101, 154))
    
    # Row 3: Doctor, Unit
    cell = info_table.cell(2, 0)
    add_left_text(cell.paragraphs[0], f"Bác sĩ chỉ định: {labtest.patient.referring_doctor or ''}", size=12, color=(84, 101, 154))
    cell = info_table.cell(2, 1)
    add_left_text(cell.paragraphs[0], "Đơn vị: MEDILAB SÀI GÒN - CẦN THƠ [00.035]", size=12, color=(84, 101, 154))

    # Make all borders white
    whiten_table_borders(info_table)

    # Add space before results table
    doc.add_paragraph()

    # Results Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'  # Keep grid style for results table
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.625)
    table.columns[1].width = Inches(1.625)
    table.columns[2].width = Inches(1.625)
    table.columns[3].width = Inches(1.625)
    
    # Header của bảng
    header_cells = table.rows[0].cells
    headers = ['XÉT NGHIỆM', 'KẾT QUẢ', 'KHOẢNG THAM CHIẾU', 'ĐƠN VỊ']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Thêm category name
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[3])
    row_cells[0].text = labtest.category.name
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in row_cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Thêm kết quả - bao gồm cả các trường không có kết quả
    for item in labtest.category.items.all().order_by('order'):
        row_cells = table.add_row().cells
        row_cells[0].text = item.name
        # Tìm kết quả tương ứng nếu có
        result = labtest.results.filter(item=item).first()
        row_cells[1].text = result.value if result else ''
        row_cells[2].text = item.reference_range or ''
        row_cells[3].text = item.unit or ''
        
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Add space before signature block
    for _ in range(3):  # Add 3 blank paragraphs for spacing
        doc.add_paragraph()

    # Signature Block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, f"TP. Cần Thơ, ngày {datetime.now().strftime('%d')} tháng {datetime.now().strftime('%m')} năm {datetime.now().strftime('%Y')}", size=12, color=(84, 101, 154))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, "PHÒNG XÉT NGHIỆM", size=12, color=(84, 101, 154))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_right_text(p, "(Ký, đóng dấu và ghi rõ họ tên)", size=12, color=(84, 101, 154))
    
    # Add signature image if exists
    if os.path.exists('static/images/signature.png'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture('static/images/signature.png', width=Cm(3))

    # Lưu file
    filename = f"labtest_{labtest.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    # Update last print date
    labtest.last_print_date = datetime.now()
    labtest.save()
    
    return filename 